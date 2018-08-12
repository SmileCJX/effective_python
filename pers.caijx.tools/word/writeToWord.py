#! python3
# writeToWord.docx - 写入word文档
import docx
doc = docx.Document()
doc.add_paragraph('hello_world!')
doc.save('.\\helloworld.docx')
paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run('This text is being added to the second paragraph.')
doc.save('.\\multiParagraphs.docx')