#! python3
# changeDocxStyle.py - Run属性
# todo hava some fault
import docx
doc = docx.Document('.\\demo.docx')
print(doc.paragraphs[0].text)
print(doc.paragraphs[0].style)
doc.paragraphs[0].style = 'Normal'
print(doc.paragraphs[1].text)
(doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)
doc.paragraphs[1].runs[0].style = 'Quotechar'
doc.paragraphs[1].runs[1].underline = 'True'
doc.paragraphs[1].runs[3].underline = 'True'
doc.save('.\\restyled.docx')
